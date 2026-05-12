# Copyright project_phoenix
"""
Legal Document Processor Module
Handles loading and parsing legal documents (PDF, DOCX, TXT).
"""

import os
import re
import zipfile
from PyPDF2 import PdfReader
from docx import Document


def _normalize_extracted_text(text):
    """Normalize extracted text without flattening paragraph structure."""
    normalized = text.replace("\x00", "")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+\n", "\n", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()



def _split_sentences(text):
    """Split text into sentences (basic sentence tokenizer)."""
    # Simple sentence tokenizer — handles common legal abbreviations
    text = re.sub(r'\b(v|vs|vs\.|et al|etc|e\.g|i\.e|esp|approx|ave|dept|inc|jr|lr|no|pp|rev|sr|vol)\.\s', 
                  r'\1_DOTMARKER ', text)
    
    sentences = re.split(r'(?<=[.!?])\s+', text)
    
    # Restore the abbreviations
    sentences = [s.replace("_DOTMARKER", ".") for s in sentences]
    
    return [s.strip() for s in sentences if s.strip()]


def _split_paragraphs(text):
    """Split text into paragraphs."""
    paragraphs = re.split(r"\n\s*\n", text)
    return [p.strip() for p in paragraphs if p.strip()]


def _extract_text(file_path):
    """Extract text from plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        raise RuntimeError(f"Error reading text file: {e}")


def _extract_docx(file_path):
    """Extract text from DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        raise RuntimeError(f"Error reading DOCX: {e}")
    
    return text


def _extract_rtf(file_path):
    """Extract readable text from a basic RTF document."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            raw = f.read()
    except Exception as e:
        raise RuntimeError(f"Error reading RTF: {e}")

    def _decode_hex(match):
        try:
            return bytes.fromhex(match.group(1)).decode("cp1252", errors="replace")
        except Exception:
            return ""

    text = re.sub(r"\\'([0-9a-fA-F]{2})", _decode_hex, raw)
    text = text.replace("\\par", "\n").replace("\\line", "\n").replace("\\tab", "\t")
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)
    text = re.sub(r"[{}]", "", text)
    text = text.replace("\\\\", "\\").replace("\\{", "{").replace("\\}", "}")
    return text


class DocumentProcessor:
    """Process and extract text from legal documents."""

    SUPPORTED_FORMATS = [".pdf", ".docx", ".txt", ".rtf", ".md"]

    def __init__(self):
        self.raw_text = ""
        self.paragraphs = []
        self.sentences = []
        self.metadata = {}

    def _detect_format(self, file_path):
        """Detect document format from extension or file signature."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext in self.SUPPORTED_FORMATS:
            return ext

        try:
            with open(file_path, "rb") as f:
                header = f.read(4096)
        except Exception as e:
            raise RuntimeError(f"Error inspecting file format: {e}")

        if header.startswith(b"%PDF-"):
            return ".pdf"

        if header.startswith(b"{\\rtf"):
            return ".rtf"

        if header.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
            try:
                with zipfile.ZipFile(file_path) as archive:
                    if "word/document.xml" in archive.namelist():
                        return ".docx"
            except zipfile.BadZipFile:
                pass

        for encoding in ("utf-8", "cp1252"):
            try:
                header.decode(encoding)
                return ".txt"
            except UnicodeDecodeError:
                continue

        return ext

    def load_document(self, file_path):
        """
        Load a document from file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            dict: Document content and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = self._detect_format(file_path)

        if ext not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported format: {ext}. Supported: {self.SUPPORTED_FORMATS}")

        if ext == ".pdf":
            text = self._extract_pdf(file_path)
        elif ext == ".docx":
            text = _extract_docx(file_path)
        elif ext == ".rtf":
            text = _extract_rtf(file_path)
        elif ext in [".txt", ".md"]:
            text = _extract_text(file_path)
        else:
            raise ValueError(f"Cannot process format: {ext}")

        text = _normalize_extracted_text(text)
        self.raw_text = text
        self.paragraphs = _split_paragraphs(text)
        self.sentences = _split_sentences(text)

        self.metadata = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "format": ext,
            "char_count": len(text),
            "word_count": len(text.split()),
            "paragraph_count": len(self.paragraphs),
            "sentence_count": len(self.sentences),
        }

        return {
            "text": text,
            "metadata": self.metadata,
            "paragraphs": self.paragraphs,
            "sentences": self.sentences,
        }

    def _extract_pdf(self, file_path):
        """Extract text from PDF file."""
        text = ""
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # Store PDF metadata
            if reader.metadata:
                self.metadata["title"] = reader.metadata.get("title", "")
                self.metadata["author"] = reader.metadata.get("author", "")
                
        except Exception as e:
            raise RuntimeError(f"Error reading PDF: {e}")
        
        return text

    def find_legal_citations(self, text):
        """
        Find legal citations in text (Supreme Court, appellate courts, etc.)
        
        Returns:
            list: Found citations with details
        """
        citations = []

        # Supreme Court of Canada
        scc_pattern = re.compile(
            r'([A-Z][a-zA-Z\s\.]+)\s+v\.?\s+([A-Z][a-zA-Z\s\.]+),?\s*'
            r'(\d{4})\s+(SCC|FC|FCA|ONCA|ONSC|BCCA|BCSC|ABCA|ABQB|QCCA|QCCS)',
            re.MULTILINE
        )

        # Neutral citation pattern
        neutral_pattern = re.compile(
            r'(\d{4})\s+(SCC|FC|FCA|ONCA|ONSC|BCCA|BCSC|ABCA|ABQB|QCCA|QCCS)\s+(\d+)',
            re.MULTILINE
        )

        # Legislation references
        legislation_pattern = re.compile(
            r'(Criminal Code|Controlled Drugs and Substances Act|Charter of Rights and Freedoms|'
            r'Youth Criminal Justice Act|Food and Drugs Act|Cannabis Act|'
            r'Firearms Act|Corrections and Conditional Release Act|'
            r'Controlled Drugs and Substances Act|Narcotic Control Act)[,\s]+'
            r'(R\.?S\.?C\.?\s*\d{4},?\s*c\.?\s*\d+|S\.?C\.?\s*\d{4},?\s*c\.?\s*\d+|SC\s*\d{4},?\s*c\.?\s*\d+)?'
            r'(?:\s*[\s,]\s*(?:s(?:ection)?\.?\s*\d+(?:\.\d+)?(?:\s*[-–]\s*\d+(?:\.\d+)?)?)?)?',
            re.IGNORECASE
        )

        # Section references
        section_pattern = re.compile(
            r'(?:s(?:ection)?\.?\s*)(\d+(?:\.\d+)?(?:\([a-z0-9]+\))?(?:\s*[-–]\s*\d+(?:\.\d+)?(?:\([a-z0-9]+\))?)?)',
            re.IGNORECASE
        )

        # Charter section references
        charter_pattern = re.compile(
            r'(?:Charter|Canadian Charter of Rights and Freedoms)[,\s]+'
            r'(?:s(?:ection)?\.?\s*)(\d+(?:\([a-z0-9]+\))?)',
            re.IGNORECASE
        )

        # Also catch "section X of the Charter" 
        charter_pattern2 = re.compile(
            r's(?:ection)?\.?\s*(\d+(?:\([a-z0-9]+\))?)\s+of\s+(?:the\s+)?(?:Charter|Canadian Charter)',
            re.IGNORECASE
        )

        # Find all citations
        for match in scc_pattern.finditer(text):
            citations.append({
                "type": "case_law",
                "full_citation": match.group(0).strip(),
                "plaintiff": match.group(1).strip(),
                "defendant": match.group(2).strip(),
                "year": match.group(3),
                "court": match.group(4),
            })

        for match in neutral_pattern.finditer(text):
            citations.append({
                "type": "neutral_citation",
                "full_citation": match.group(0).strip(),
                "year": match.group(1),
                "court": match.group(2),
                "number": match.group(3),
            })

        for match in legislation_pattern.finditer(text):
            citations.append({
                "type": "legislation",
                "full_citation": match.group(0).strip(),
                "statute": match.group(1),
            })

        for match in charter_pattern.finditer(text):
            citations.append({
                "type": "charter_section",
                "full_citation": match.group(0).strip(),
                "section": match.group(1),
            })

        for match in charter_pattern2.finditer(text):
            citations.append({
                "type": "charter_section",
                "full_citation": match.group(0).strip(),
                "section": match.group(1),
            })

        # Deduplicate
        seen = set()
        unique_citations = []
        for c in citations:
            key = c.get("full_citation", "")
            if key not in seen:
                seen.add(key)
                unique_citations.append(c)

        return unique_citations

    def find_key_legal_terms(self, text):
        """Find key legal terms and phrases in the document."""
        legal_terms = {
            "charter_breach": [
                "charter breach", "constitutional violation", "constitutional breach",
                "rights violation", "infringement of rights", "charter infringement",
                "charter violation"
            ],
            "remedies": [
                "section 24(1)", "s. 24(1)", "exclusion of evidence", "s. 24(2)",
                "stay of proceedings", "damages", "constitutional remedy",
                "grant charter remedy", "dismiss charter application"
            ],
            "procedures": [
                "voir dire", "preliminary inquiry", "trial within a reasonable time",
                "disclosure", "adjournment", "continuation", "directed verdict",
                "no case to answer", "motion to exclude", "charter application",
                "within a reasonable time", "judicial stay"
            ],
            "evidence": [
                "admissibility", "excluded", "confession", "statement",
                "voluntariness", "hearsay", "similar fact evidence", "character evidence",
                "illegally obtained", "derivative evidence", "fruit of the poisonous tree",
                "good faith", "consent search", "warrantless search"
            ],
            "standards": [
                "beyond reasonable doubt", "balance of probabilities", "preponderance of evidence",
                "reasonable grounds", "probable cause", "reasonable suspicion",
                "prima facie", "onus of proof", "reverse onus", "burden of proof"
            ],
            "detention_arrest": [
                "arrest", "detention", "investigative detention", "arbitrary detention",
                "right to counsel", "duty to inform", "promptly informed",
                "right to silence", "right to retain counsel"
            ],
            "sentencing": [
                "mandatory minimum", "proportionality", "totality", "parity",
                "deterrence", "rehabilitation", "denunciation", "restitution",
                "conditional sentence", "probation", "intermittent sentence",
                "credit for pre-trial custody", "fine"
            ],
        }

        found_terms = {}
        text_lower = text.lower()

        for category, terms in legal_terms.items():
            found = []
            for term in terms:
                count = text_lower.count(term.lower())
                if count > 0:
                    found.append({"term": term, "count": count})
            if found:
                found_terms[category] = sorted(found, key=lambda x: x["count"], reverse=True)

        return found_terms

    def get_statistics(self):
        """Get document statistics."""
        return {
            "char_count": len(self.raw_text),
            "word_count": len(self.raw_text.split()),
            "paragraph_count": len(self.paragraphs),
            "sentence_count": len(self.sentences),
            "avg_sentence_length": len(self.raw_text.split()) / max(len(self.sentences), 1),
            "avg_paragraph_length": len(self.raw_text.split()) / max(len(self.paragraphs), 1),
        }

    def extract_all_text(self, file_path):
        """Comprehensive extraction function designed for Android load integration."""
        try:
            return self.load_document(file_path)
        except Exception as e:
            return {"text": "", "error": str(e), "metadata": {}}
